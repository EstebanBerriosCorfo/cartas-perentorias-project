import os
from datetime import datetime
from docx import Document
from docx.shared import Pt

"""
core/letter_generator.py
Genera cartas perentorias o de incumplimiento en formato Word,
utilizando los datos obtenidos desde la interfaz.
"""

class LetterGenerator:
    """Generador básico de cartas perentorias e incumplimiento."""

    def __init__(self):
        # Carpeta de salida (Descargas del usuario)
        self.output_folder = os.path.join(os.path.expanduser("~/Downloads"), "Cartas Generadas")
        os.makedirs(self.output_folder, exist_ok=True)

    # ─────────────────────────────────────────────
    # 🔹 MÉTODO PRINCIPAL
    # ─────────────────────────────────────────────
    def generar_carta(self, codigo_proyecto, tipo_carta, informe_asociado, datos_proyecto=None):
        """
        Genera un documento Word en base al tipo de carta y los datos del proyecto.
        :param codigo_proyecto: str
        :param tipo_carta: str ("Generar carta perentoria" o "Generar carta de incumplimiento")
        :param informe_asociado: str (nombre del informe seleccionado)
        :param datos_proyecto: dict (opcional: con nombreProyecto, beneficiario, responsable)
        :return: ruta del archivo generado
        """

        # Crear documento base
        doc = Document()

        # Encabezado
        doc.add_heading("Carta " + tipo_carta.replace("Generar ", ""), level=1)

        # Fecha
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        doc.add_paragraph(f"Santiago, {fecha_actual}")

        # Cuerpo del documento
        doc.add_paragraph(f"Código del proyecto: {codigo_proyecto}")

        if datos_proyecto:
            doc.add_paragraph(f"Nombre del Proyecto: {datos_proyecto.get('nombreProyecto', '')}")
            doc.add_paragraph(f"Beneficiario: {datos_proyecto.get('beneficiario', '')}")
            doc.add_paragraph(f"Representante Legal: {datos_proyecto.get('representanteLegal', '')}")

        doc.add_paragraph(f"Informe asociado: {informe_asociado}")

        # Cuerpo principal según tipo de carta
        cuerpo = (
            "Por medio de la presente, se notifica al beneficiario que el informe mencionado "
            "se encuentra pendiente de entrega dentro de los plazos establecidos por CORFO."
        )

        if "incumplimiento" in tipo_carta.lower():
            cuerpo = (
                "De acuerdo con los antecedentes revisados, se constata un incumplimiento "
                "en la entrega del informe indicado, conforme a lo establecido en la resolución vigente."
            )

        doc.add_paragraph(cuerpo)

        doc.add_paragraph("\nSin otro particular,")
        doc.add_paragraph("Atentamente,")
        doc.add_paragraph("Subdirección de Operaciones y Mejora Continua")
        doc.add_paragraph("Gerencia de Innovación – CORFO")

        # Formato básico
        for p in doc.paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        # Guardar archivo
        nombre_archivo = f"{tipo_carta.replace('Generar ', '').replace(' ', '_')}_{codigo_proyecto}.docx"
        ruta_salida = os.path.join(self.output_folder, nombre_archivo)
        doc.save(ruta_salida)

        return ruta_salida