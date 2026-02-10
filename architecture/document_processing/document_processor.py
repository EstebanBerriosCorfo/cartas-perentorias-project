import os
from datetime import datetime
from docx import Document
from architecture.utils.path_utils import generate_download_path

# Mapa de meses en español (evitamos depender del locale del sistema)
SPANISH_MONTHS = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

class DocumentProcessor:
    """
    Genera cartas (Perentoria / Incumplimiento) desde plantillas Word
    y datos integrados (dict).
    """

    def __init__(self):
        self.template_dir = os.path.join(os.path.dirname(__file__), "..", "document_templates")

    # -----------------------------
    # Util
    # -----------------------------
    def _get_template_path(self, letter_type: str) -> str:
        templates = {
            "perentoria": "Carta_Perentoria.docx",
            "incumplimiento": "Carta_Incumplimiento_Informe.docx"
        }
        file_name = templates.get(letter_type.lower())
        if not file_name:
            raise ValueError(f"Tipo de carta no reconocido: {letter_type}")
        return os.path.join(self.template_dir, file_name)

    def _fmt_fecha(self, fecha: datetime) -> tuple[str, str, int]:
        """Devuelve (día, mes_en_español, año)"""
        return str(fecha.day), SPANISH_MONTHS[fecha.month], fecha.year

    def _parse_date_for_sort(self, value: str) -> datetime:
        if not value:
            return datetime.max
        try:
            return datetime.strptime(str(value).strip(), "%d/%m/%Y")
        except Exception:
            return datetime.max

    def _build_tipo_informe(self, reports: list, report: dict) -> str:
        report_type = str(report.get("reportType", "")).strip()
        if not report_type:
            return ""

        def _norm_tipo(r):
            return str(r.get("reportType", "")).strip().upper()

        same_type = [r for r in reports if _norm_tipo(r) == report_type.upper()]
        if len(same_type) <= 1:
            return report_type

        same_type_sorted = sorted(
            same_type,
            key=lambda r: self._parse_date_for_sort(r.get("scheduledDeliveryDate", ""))
        )
        sel_date = str(report.get("scheduledDeliveryDate", "")).strip()
        index = None
        for i, r in enumerate(same_type_sorted, start=1):
            if str(r.get("scheduledDeliveryDate", "")).strip() == sel_date:
                index = i
                break
        if index is None:
            index = 1

        return f"{report_type} {index}"

    def _replace_everywhere(self, doc: Document, replacements: dict):
        """
        Reemplazo robusto que funciona aunque el marcador esté fragmentado en runs.
        Nota: reescribe el texto del párrafo/celda (se puede perder formato dentro del marcador).
        """

        # Párrafos
        for p in doc.paragraphs:
            original = p.text
            new_text = original
            for k, v in replacements.items():
                if k in new_text:
                    new_text = new_text.replace(k, str(v))
            if new_text != original:
                # Limpia runs y deja un solo run con el texto reemplazado
                for r in p.runs:
                    r.clear()  # limpia contenido del run
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    p.add_run(new_text)

        # Tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original = cell.text
                    new_text = original
                    for k, v in replacements.items():
                        if k in new_text:
                            new_text = new_text.replace(k, str(v))
                    if new_text != original:
                        # Reemplazo básico: limpiar y escribir plano
                        # (Word vuelve a componer internamente los párrafos)
                        for para in cell.paragraphs:
                            for r in para.runs:
                                r.clear()
                        if cell.paragraphs:
                            cell.paragraphs[0].runs[0].text = new_text
                        else:
                            cell.add_paragraph(new_text)

    # -----------------------------
    # Público
    # -----------------------------
    def generate_letter(self, data: dict, report_type: str, report_date: str | None, letter_type: str) -> str:
        # 1) Selección de informe
        reports = data.get("reports", [])
        print("🔍 report_type recibido:", report_type)
        print("🗓️ report_date recibido:", report_date)
        print("📄 tipos disponibles:", [r.get("reportType") for r in reports])
        if report_date:
            report = next(
                (
                    r for r in reports
                    if r.get("reportType", "").strip().upper() == report_type.strip().upper()
                    and str(r.get("scheduledDeliveryDate", "")).strip() == str(report_date).strip()
                ),
                None
            )
        else:
            report = next(
                (
                    r for r in reports
                    if r.get("reportType", "").strip().upper() == report_type.strip().upper()
                ),
                None
            )
        if not report:
            detalle_fecha = f" con fecha {report_date}" if report_date else ""
            raise ValueError(f"No se encontró el informe '{report_type}'{detalle_fecha} en los datos del proyecto.")

        # 2) Carga de plantilla
        template_path = self._get_template_path(letter_type)
        doc = Document(template_path)

        # 3) Datos
        project = data["projectinfo"]

        # Fechas (informe y resolución)
        fecha_entrega = datetime.strptime(report["scheduledDeliveryDate"], "%d/%m/%Y")
        fecha_resol   = datetime.strptime(project["resolutionDate"], "%d/%m/%Y")

        dia_inf, mes_inf, anio_inf = self._fmt_fecha(fecha_entrega)
        dia_res, mes_res, anio_res = self._fmt_fecha(fecha_resol)

        # 🔍 Lógica jerárquica para determinar el correo de contacto
        direccion = (
            project.get("legalRepresentativeEmail")
            or project.get("beneficiaryEmail")
            or project.get("directorEmail")
            or "SIN CORREO REGISTRADO"
        )
        direccion = direccion.strip() if isinstance(direccion, str) else "SIN CORREO REGISTRADO"

        # 4) Replacements
        tipo_informe = self._build_tipo_informe(reports, report)
        replacements = {
            # Identificación
            "[NOMBRE INFORME]": report["reportType"],
            "[TIPO INFORME]": tipo_informe,
            "[NOMBRE DE PROYECTO]": project["projectName"].strip(),
            "[CÓDIGO]": project["projectCode"],

            # Destinatario
            "[NOMBRE BENEFICIARIA]": project["beneficiaryName"].strip(),
            "[nombre representante]": project["legalRepresentative"].strip(),  # si la plantilla lo usa
            "[DIRECCIÓN]": direccion,

            # Fechas del informe
            "[DÍA]": dia_inf,
            "[MES]": mes_inf,
            "[AÑO]": anio_inf,

            # Fechas de la resolución
            "[DÍA RESOL]": dia_res,
            "[MES RESOL]": mes_res,
            "[AÑO RESOL]": anio_res,

            # Resolución y firmas
            "[NÚMERO]": int(project["resolutionNumber"]),
            "[SUBDIRECTOR]": project.get("subdirector", "").strip(),
            "[SUBDIRECCION]": project.get("subdirection", "").strip(),
            "[EJECUTIVO TÉCNICO]": project.get("technicalExecutiveName", "").strip()
        }

        # 5) Reemplazo robusto
        self._replace_everywhere(doc, replacements)

        # 6) Exportación
        output_path = generate_download_path(project["projectCode"], letter_type)
        doc.save(output_path)
        print(f"✅ Carta generada exitosamente: {output_path}")
        return output_path
